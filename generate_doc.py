"""
Generates CIT_CyberSecurity_CourseAdvisor_Mindmap.docx
Run:  python3 generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ───────────────────────────────────────────────
C_BLACK      = RGBColor(0x0f, 0x11, 0x17)
C_DARKBG     = RGBColor(0x1e, 0x29, 0x3b)
C_HEADING    = RGBColor(0x1e, 0x40, 0x8f)   # deep navy
C_SUBHEAD    = RGBColor(0x0f, 0x4c, 0x75)   # medium blue
C_BODY       = RGBColor(0x1e, 0x29, 0x3b)
C_MUTED      = RGBColor(0x47, 0x55, 0x69)
C_GREEN      = RGBColor(0x16, 0xa3, 0x4a)
C_RED        = RGBColor(0xdc, 0x26, 0x26)
C_ORANGE     = RGBColor(0xd9, 0x77, 0x06)
C_BLUE       = RGBColor(0x1d, 0x4e, 0xd8)
C_PURPLE     = RGBColor(0x5b, 0x21, 0xb6)
C_TEAL       = RGBColor(0x0f, 0x76, 0x6e)

# shading fills
SH_NAVY      = "1e3a5f"
SH_GREY      = "f1f5f9"
SH_LIGHT     = "e0f2fe"
SH_GREEN     = "dcfce7"
SH_RED       = "fee2e2"
SH_YELLOW    = "fef9c3"
SH_ORANGE    = "ffedd5"
SH_PURPLE    = "ede9fe"
SH_TEAL      = "ccfbf1"
SH_WHITE     = "FFFFFF"

doc = Document()

# ── Page margins ─────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Helpers ──────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_borders(cell, color="CCCCCC", sz=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def run(para, text, bold=False, italic=False, color=None, size=None, underline=False):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.underline = underline
    if color:  r.font.color.rgb = color
    if size:   r.font.size = Pt(size)
    return r

def heading(text, level=1, color=C_HEADING, size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color
    if level == 1:
        r.font.size = Pt(size or 20)
    elif level == 2:
        r.font.size = Pt(size or 14)
    else:
        r.font.size = Pt(size or 11)
    pPr = p._p.get_or_add_pPr()
    # spacing
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), "120" if level == 1 else "200")
    pSp.set(qn("w:after"),  "80"  if level == 1 else "60")
    pPr.append(pSp)
    return p

def section_label(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = C_MUTED
    r.font.all_caps = True
    pPr = p._p.get_or_add_pPr()
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), "280")
    pSp.set(qn("w:after"),  "60")
    pPr.append(pSp)
    # bottom border
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CBD5E1")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def body(text, color=None, size=10, space_before=0, space_after=60, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color: r.font.color.rgb = color
    pPr = p._p.get_or_add_pPr()
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), str(space_before))
    pSp.set(qn("w:after"),  str(space_after))
    pPr.append(pSp)
    return p

def bullet(text, indent=1, color=None, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    pPr = p._p.get_or_add_pPr()
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), "20")
    pSp.set(qn("w:after"),  "20")
    pPr.append(pSp)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),    str(360 * indent))
    ind.set(qn("w:hanging"), "180")
    pPr.append(ind)
    return p

def note_box(text, fill=SH_YELLOW, text_color=None, icon="", size=9):
    """Single-cell 1×1 table used as a shaded note box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0,0)
    set_cell_bg(cell, fill)
    cell_borders(cell, color="DDDDDD", sz=6)
    p = cell.paragraphs[0]
    if icon:
        run(p, icon + "  ", size=size)
    run(p, text, color=text_color, size=size)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    # spacer after box
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def hline():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CBD5E1")
    pBdr.append(bot)
    pPr.append(pBdr)
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), "160")
    pSp.set(qn("w:after"),  "0")
    pPr.append(pSp)

def add_page_break():
    doc.add_page_break()

# ────────────────────────────────────────────────────────────────
#  COVER / TITLE
# ────────────────────────────────────────────────────────────────
heading("CIT Cyber Security — Course Advisor", level=1, color=C_HEADING,
        size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, "Logistics Mindmap: Programs · Catalog Maps · Course Rules · Special Cases",
    color=C_MUTED, size=10)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p2, "Reference document for instructors & advisors  ·  Not for student distribution",
    color=C_MUTED, size=8.5, italic=True)

hline()

# ────────────────────────────────────────────────────────────────
#  SECTION 1 — STUDENT DECISION FLOW
# ────────────────────────────────────────────────────────────────
section_label("1 — Student Decision Flow")

# Flow table: 4 steps + outcomes
t = doc.add_table(rows=2, cols=5)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.style = "Table Grid"

headers = ["Step 1", "Step 2", "Step 3", "Step 4", "Possible Outcomes"]
titles  = ["Select Program", "Select Catalog Year", "Complete Checklist", "View Results", ""]
subs    = [
    "Compliance AAS\nDF AAS\nDF CA\nNet-Sec AAS",
    "Which degree map am I following?\n(2025–26 / 2026–27 / 2027–28)",
    "Check off every course completed or currently in progress",
    "See remaining courses, warnings, compare option",
    "✅ Graduation-ready\n📊 Compare to new map\n⚠️ Substitution needed\n👤 See advisor",
]
fills = [SH_LIGHT, SH_LIGHT, SH_LIGHT, SH_LIGHT, SH_GREEN]

for col_i, (hdr, title, sub, fill) in enumerate(zip(headers, titles, subs, fills)):
    cell = t.cell(0, col_i)
    set_cell_bg(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, hdr + "\n", bold=True, color=C_BLUE, size=7.5)
    if title:
        run(p, title, bold=True, color=C_BODY, size=9.5)

    cell2 = t.cell(1, col_i)
    set_cell_bg(cell2, SH_WHITE)
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, sub, color=C_MUTED, size=8)

# column widths
widths = [1.5, 1.8, 2.0, 1.8, 2.0]
for row in t.rows:
    for i, cell in enumerate(row.cells):
        cell.width = Inches(widths[i])

doc.add_paragraph()

# ────────────────────────────────────────────────────────────────
#  SECTION 2 — PROGRAMS
# ────────────────────────────────────────────────────────────────
hline()
section_label("2 — Programs & Their Course Structures")

# ── helper to build a program block ──────────────────────────────
def program_block(icon, title, credential, maps, core_courses,
                  elective_info, digital_lit, old_courses,
                  equivalencies, warnings):
    # Title row
    p = doc.add_paragraph()
    run(p, f"{icon}  {title}", bold=True, color=C_HEADING, size=12)
    run(p, f"  —  {credential}", color=C_MUTED, size=9)
    pPr = p._p.get_or_add_pPr()
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), "200")
    pSp.set(qn("w:after"),  "40")
    pPr.append(pSp)

    # Map version pills inline
    p2 = doc.add_paragraph()
    for i, m in enumerate(maps):
        colors_map = {"2025–26": (SH_GREY, C_MUTED),
                      "2026–27": (SH_LIGHT, C_BLUE),
                      "2027–28": (SH_PURPLE, C_PURPLE)}
        fill, fc = colors_map.get(m, (SH_GREY, C_MUTED))
        run(p2, f"  {m}  ", bold=True, color=fc, size=8.5)
    p2.paragraph_format.space_after = Pt(4)

    # details table
    rows_data = []
    if core_courses:
        rows_data.append(("Core Courses", core_courses, SH_LIGHT, C_BLUE))
    if elective_info:
        rows_data.append(("Electives / Tracks", elective_info, SH_TEAL, C_TEAL))
    if digital_lit:
        rows_data.append(("Digital Literacy", digital_lit, SH_PURPLE, C_PURPLE))
    if old_courses:
        rows_data.append(("Old Map ⚠️", old_courses, SH_ORANGE, C_ORANGE))
    if equivalencies:
        rows_data.append(("Course Equiv.", equivalencies, SH_GREEN, C_GREEN))

    t = doc.add_table(rows=len(rows_data), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    col_widths = [Inches(1.4), Inches(7.2)]
    for ri, (label, content, fill, lc) in enumerate(rows_data):
        lc_cell = t.cell(ri, 0)
        vc_cell = t.cell(ri, 1)
        set_cell_bg(lc_cell, "E2E8F0")
        set_cell_bg(vc_cell, SH_WHITE)
        lc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lc_cell.paragraphs[0]
        run(lp, label, bold=True, color=C_SUBHEAD, size=8.5)
        vp = vc_cell.paragraphs[0]
        run(vp, content, color=C_BODY, size=9)
        for cell in (lc_cell, vc_cell):
            cell_borders(cell, "DDDDDD", 4)
        lc_cell.width = col_widths[0]
        vc_cell.width = col_widths[1]

    # warnings
    for icon_w, text_w, fill_w in warnings:
        note_box(text_w, fill=fill_w, icon=icon_w, size=8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# Compliance AAS
program_block(
    icon="📋", title="Cyber Security — Compliance", credential="AAS Degree",
    maps=["2025–26", "2026–27"],
    core_courses="⚠️ PLACEHOLDER — course data not yet entered. Contact instructor.",
    elective_info="",
    digital_lit="IS 100B (0 credits)  ·  IS 101 (3 credits)",
    old_courses="",
    equivalencies="",
    warnings=[
        ("⚠️", "All Compliance AAS course data is placeholder. Awaiting instructor input before this program is usable.", SH_YELLOW),
    ]
)

# DF AAS
program_block(
    icon="🔍", title="Cyber Security — Digital Forensics", credential="AAS Degree",
    maps=["2025–26", "2026–27", "2027–28"],
    core_courses=(
        "2026–27 (9 courses):  CIT 114 · CIT 173 · CIT 217 · CSEC 101 · CSEC 110 · "
        "CSEC 120 · CSEC 121 · CSEC 191 · CSEC 220\n"
        "2027–28 adds (required core):  CIT 112 · CSEC 221"
    ),
    elective_info=(
        "2026–27: 6–9 elective credits depending on Digital Literacy choice\n"
        "Recommended: CSEC 221 (highly rec.) · CIT 112 (highly rec.) · Any CIT/CSEC course\n"
        "2027–28: No electives — fully prescribed (11 core courses)"
    ),
    digital_lit=(
        "IS 100B (0 cr) → requires 9 elective credits  ·  "
        "IS 101 (3 cr) → requires 6 elective credits"
    ),
    old_courses=(
        "CSEC 112B — DEACTIVATED (no longer offered)\n"
        "CSEC 116B — ON HOLD (not currently offered)\n"
        "CSEC 212B — ON HOLD (not currently offered)\n"
        "CSEC 290B — LAST OFFERED: Fall 2026 (AAS only — enroll immediately if still needed)\n"
        "CIT 263 — Required in old map only · Does NOT carry to new map"
    ),
    equivalencies="CSEC 110B = CSEC 110 (direct code rename)",
    warnings=[
        ("⚠️", "CSEC 290B is AAS-only. CA students must NOT enroll. Last offering: Fall 2026.", SH_ORANGE),
        ("ℹ️", "CIT 263 (Project Management) does NOT carry to the 2026–27 or 2027–28 map.", SH_LIGHT),
    ]
)

# DF CA
program_block(
    icon="📂", title="Cyber Security — Digital Forensics", credential="CA — Certificate of Achievement",
    maps=["2025–26", "2026–27"],
    core_courses=(
        "2026–27 (10 courses, fully prescribed):  CIT 112 · CIT 114 · CIT 217 · CSEC 101 · "
        "CSEC 110 · CSEC 120 · CSEC 121 · CSEC 191 · CSEC 220 · CSEC 221"
    ),
    elective_info="None — the CA is fully prescribed. No open electives.",
    digital_lit="IS 100B (0 credits)  ·  IS 101 (3 credits)",
    old_courses=(
        "CSEC 112B — DEACTIVATED · CSEC 116B — ON HOLD · CSEC 212B — ON HOLD\n"
        "CIT 211 — Required in old CA only · Does NOT carry to new CA\n"
        "CIT 263 — Required in old CA only · Does NOT carry to new CA"
    ),
    equivalencies="CSEC 110B = CSEC 110 (direct code rename)",
    warnings=[
        ("⛔", "CSEC 290B is NOT part of the CA. CA students must NEVER enroll in CSEC 290B.", SH_RED),
        ("ℹ️", "CIT 211 and CIT 263 (old required courses) do NOT carry to the new CA map.", SH_LIGHT),
    ]
)

# Net-Sec AAS
program_block(
    icon="🌐", title="Cyber Security — Network Security", credential="AAS Degree",
    maps=["2025–26", "2026–27", "2027–28"],
    core_courses=(
        "CIT 105 · CIT 112 · CIT 114 · CIT 126 · CIT 129 · CIT 173 · CIT 217 · "
        "CSEC 286 (take last, after all other core)"
    ),
    elective_info=(
        "Choose ONE of 6 tracks (12 credits each):\n"
        "Cisco · Cloud Security · Digital Forensics · Compliance · Offensive Security · Artificial Intelligence\n"
        "2027–28: OffSec adds CSEC 282 · AI track adds CSEC 104 + CSEC 282 (both become full 4-course tracks)"
    ),
    digital_lit="IS 100B (0 credits)  ·  IS 101 (3 credits)",
    old_courses=(
        "CSEC 114B — DEACTIVATED · CSCO 105B — Now ET-only (not available for CIT students)"
    ),
    equivalencies=(
        "CSCO 230B = CSCO 230 (Cisco track 3rd slot)  ·  "
        "CSEC 281B = CIT 274 (OffSec track, course 1) — NOT the same as CSEC 281  ·  "
        "CSEC 104B = CSEC 104 (AI track, 2027–28 only)"
    ),
    warnings=[
        ("✅", "No deactivated courses in the new Net-Sec maps — clean transition for most students.", SH_GREEN),
        ("⚠️", "Old map: CSEC 114B deactivated. CSCO 105B now ET-only. CSEC 290B is NOT for Net-Sec students.", SH_ORANGE),
    ]
)

# ────────────────────────────────────────────────────────────────
#  SECTION 3 — SPECIAL RULES
# ────────────────────────────────────────────────────────────────
add_page_break()
section_label("3 — Special Rules & Edge Cases")

rules = [
    {
        "icon": "🔒",
        "title": "CSEC 181 → CSEC 281 Prerequisite Lock (Net-Sec, Offensive Security track)",
        "fill": SH_ORANGE,
        "body": (
            "CSEC 281 can ONLY be taken after CSEC 181 is completed with a grade of C or higher. "
            "It cannot be taken before, during, or at the same time as CSEC 181.\n\n"
            "Important distinction: CSEC 281B = CIT 274 (same course, different code). "
            "CSEC 281B is NOT the same course as CSEC 281. Students often confuse these."
        )
    },
    {
        "icon": "🔄",
        "title": "CHFI Substitution Sequence (DF AAS & CA old-map students)",
        "fill": SH_GREEN,
        "body": (
            "If CSEC 112B / 116B / 212B are still needed (not yet completed), the recommended "
            "substitute path is the full CHFI certification prep sequence:\n\n"
            "    Step 1: CSEC 120 — CHFI 1: Digital Content Forensics\n"
            "    Step 2: CSEC 121 — CHFI 2: Internet & Network Forensics\n"
            "         ↓ Complete BOTH 100-level courses before enrolling in 200-level ↓\n"
            "    Step 3: CSEC 220 — CHFI 3: Operating Systems Forensics\n"
            "    Step 4: CSEC 221 — CHFI 4: Social Media, Mobile & Cloud Forensics\n\n"
            "These four courses are also required core in the 2026–27 map — making a map switch "
            "the cleanest path for most students who haven't yet started the discontinued courses."
        )
    },
    {
        "icon": "⚠️",
        "title": "3-Scenario Deactivated Course Logic (DF Programs — CSEC 112B, 116B, 212B)",
        "fill": SH_RED,
        "body": (
            "Scenario A — None of the three courses completed yet:\n"
            "  → Recommend CHFI sequence as substitute. Cleanest path.\n\n"
            "Scenario B — Some completed, some still needed:\n"
            "  → Two options:\n"
            "     Option A: Enroll in all 4 CHFI courses (120→121→220→221) in strict order.\n"
            "     Option B: Work individually with Professor Morningstar to find approved "
            "CIT/CSEC/CSCO substitutes.\n"
            "  → Courses already completed count as elective credits regardless of choice.\n\n"
            "Scenario C — All three already completed:\n"
            "  → All count as elective credits toward degree. Optional: take CHFI for certification."
        )
    },
    {
        "icon": "⏰",
        "title": "CSEC 290B — Last Offered Fall 2026 (DF AAS only)",
        "fill": SH_ORANGE,
        "body": (
            "AAS students only. CA students must NOT enroll — CSEC 290B is not part of the CA.\n\n"
            "If a DF AAS student has NOT yet completed CSEC 290B, they must register for Fall 2026 "
            "immediately — this is the final offering. After Fall 2026, this course will no longer "
            "be available under any circumstances.\n\n"
            "Handled separately from the deactivated 112B/116B/212B courses — different urgency "
            "and different message. The CHFI substitution path does NOT apply here."
        )
    },
    {
        "icon": "👤",
        "title": "Professor Morningstar — Individual Substitution",
        "fill": SH_YELLOW,
        "body": (
            "Available to students in Scenario B (partially completed discontinued DF courses). "
            "Student meets individually with Professor Morningstar to identify approved substitutes.\n\n"
            "Requirements:\n"
            "  • Substitutes must be CIT, CSEC, or CSCO courses\n"
            "  • Student must bring transcript and current course plan to the meeting\n"
            "  • All substitutions must be individually approved — no pre-set list"
        )
    },
    {
        "icon": "↔️",
        "title": "Course Code Equivalencies (same course, different catalog codes)",
        "fill": SH_LIGHT,
        "body": (
            "Old code       →  Current code\n"
            "CSEC 110B      →  CSEC 110   (Introduction to Digital Forensics)\n"
            "CSEC 281B      →  CIT 274    (Ethical Hacking Essentials)\n"
            "CSCO 230B      →  CSCO 230   (Security Operations Center Fundamentals)\n"
            "CSEC 104B      →  CSEC 104   (AI Security Essentials — Net-Sec 2027–28)\n"
            "CSEC 112 / CIT 112  →  same course (crosslisted)\n\n"
            "⛔ Critical: CSEC 281B ≠ CSEC 281. They are different courses. "
            "CSEC 281B is CIT 274; CSEC 281 is a CEH 2 course requiring CSEC 181 as a prerequisite."
        )
    },
    {
        "icon": "💳",
        "title": "4-Credit Course Exceptions (all others are 3 credits)",
        "fill": SH_TEAL,
        "body": (
            "The following courses are 4 credits (not the standard 3):\n"
            "CIT 114 / CSEC 114,  CSCO 105B,  CSCO 120,  CSCO 121,  CSCO 130B,\n"
            "CSCO 205B,  CSCO 220,  CSCO 230,  CSCO 230B,  CSCO 480,  CSCO 482,  CSCO 484\n\n"
            "Tuition is charged per credit. Students choosing Cisco or CSCO-heavy tracks "
            "should be aware their elective block will cost more than the standard 3-credit equivalent."
        )
    },
    {
        "icon": "📊",
        "title": "Map Comparison Feature",
        "fill": SH_PURPLE,
        "body": (
            "Old-map students can click 'See how I'd look on 2026–27 requirements →' to re-run "
            "their same completed courses against the new map requirements without re-entering data.\n\n"
            "The comparison is read-only and for reference only. To officially switch maps, "
            "the student must have an advisor appointment. The tool shows which courses carry ✅ "
            "and which don't ❌ in the new map."
        )
    },
]

for rule in rules:
    # Rule header
    p = doc.add_paragraph()
    run(p, f"{rule['icon']}  {rule['title']}", bold=True, color=C_SUBHEAD, size=10.5)
    pPr = p._p.get_or_add_pPr()
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), "180")
    pSp.set(qn("w:after"),  "40")
    pPr.append(pSp)

    note_box(rule["body"], fill=rule["fill"], size=9)


# ────────────────────────────────────────────────────────────────
#  SECTION 4 — TRANSFER TABLE
# ────────────────────────────────────────────────────────────────
add_page_break()
section_label("4 — Old-Map Courses: What Carries to the 2026–27 Map?")

xfer_data = [
    # (old_course, name, aas_status, ca_status, notes)
    ("CIT 112",   "Network+",                                 "✅ Required core",       "✅ Required core",       ""),
    ("CIT 114",   "IT Essentials",                            "✅ Required core",       "✅ Required core",       "4-credit course"),
    ("CIT 173",   "Introduction to Linux (Linux+)",           "✅ Required core",       "Not in CA core",         ""),
    ("CIT 217",   "Security+",                                "✅ Required core",       "✅ Required core",       ""),
    ("CIT 211",   "Microsoft OS Management",                  "Not in AAS",             "❌ Not in new CA",       "Old CA only — no equivalent in new maps"),
    ("CIT 263",   "Project Management",                       "❌ Not in new AAS",      "❌ Not in new CA",       "Required in old maps — does NOT carry to either new map"),
    ("CSEC 110B", "Introduction to Digital Forensics",        "✅ = CSEC 110 (core)",   "✅ = CSEC 110 (core)",   "Direct course code rename"),
    ("CSEC 112B", "Intro to Applied Windows Forensics",       "Elective credit only",   "Elective credit only",   "DEACTIVATED — counts as elective if already completed"),
    ("CSEC 116B", "Intro to Mobile Device Forensics",         "Elective credit only",   "Elective credit only",   "ON HOLD — counts as elective if already completed"),
    ("CSEC 212B", "Intermediate Applied Windows Forensics",   "Elective credit only",   "Elective credit only",   "ON HOLD — counts as elective if already completed"),
    ("CSEC 290B", "Security Capstone",                        "AAS — last Fall 2026",   "⛔ CA must NOT enroll",  "AAS-specific. No direct equivalent in new map."),
    ("CSEC 120",  "CHFI 1: Digital Content Forensics",        "✅ Required core",       "✅ Required core",       "Also used as CHFI substitute starting point"),
    ("CSEC 121",  "CHFI 2: Internet & Network Forensics",     "✅ Required core",       "✅ Required core",       ""),
    ("CSEC 191",  "Disaster Recovery & Business Continuity",  "✅ Required core",       "✅ Required core",       ""),
    ("CSEC 220",  "CHFI 3: Operating Systems Forensics",      "✅ Required core",       "✅ Required core",       ""),
    ("CSEC 221",  "CHFI 4: Social Media, Mobile & Cloud",     "✅ Rec. elective → core in 27–28", "✅ Required core (CA)", ""),
    ("CIT 126",   "Cloud Computing Fundamentals",             "Elective credit",        "General credit",         "Students pay per credit"),
    ("CIT 129",   "Introduction to Programming",              "Elective credit",        "Not in CA",              ""),
]

t = doc.add_table(rows=len(xfer_data)+1, cols=5)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.LEFT

# Header row
hdr_labels = ["Old Course", "Course Name", "AAS Carries?", "CA Carries?", "Notes"]
hdr_row = t.rows[0]
for i, lbl in enumerate(hdr_labels):
    cell = hdr_row.cells[i]
    set_cell_bg(cell, SH_NAVY.replace("#","") if SH_NAVY.startswith("#") else "1e3a5f")
    p = cell.paragraphs[0]
    run(p, lbl, bold=True, color=RGBColor(0xe2,0xe8,0xf0), size=8.5)
    cell_borders(cell, "FFFFFF", 6)

# Data rows
for ri, (course, name, aas, ca, notes) in enumerate(xfer_data):
    row = t.rows[ri+1]
    fill = "F8FAFC" if ri % 2 == 0 else "FFFFFF"

    # Determine status colors
    def status_color(text):
        if text.startswith("✅"): return C_GREEN
        if text.startswith("❌") or text.startswith("⛔"): return C_RED
        if "last" in text.lower() or "AAS —" in text: return C_ORANGE
        if "Elective" in text: return C_ORANGE
        return C_MUTED

    vals = [course, name, aas, ca, notes]
    for ci, val in enumerate(vals):
        cell = row.cells[ci]
        set_cell_bg(cell, fill)
        p = cell.paragraphs[0]
        color = status_color(val) if ci in (2,3) else (C_BLUE if ci==0 else C_BODY)
        bold  = ci == 0
        run(p, val, color=color, size=8.5, bold=bold)
        cell_borders(cell, "DDDDDD", 4)

# Column widths
col_w = [Inches(0.85), Inches(2.3), Inches(1.4), Inches(1.4), Inches(2.0)]
for row in t.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_w[i]

doc.add_paragraph()

# Legend
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run(p, "Legend:  ", bold=True, size=8.5, color=C_MUTED)
run(p, "✅ Carries as required core  ", color=C_GREEN, size=8.5)
run(p, "  Elective/general credit  ", color=C_ORANGE, size=8.5)
run(p, "  ❌ Does NOT carry  ", color=C_RED, size=8.5)
run(p, "  Grey = does not apply to that program", color=C_MUTED, size=8.5)


# ────────────────────────────────────────────────────────────────
#  SECTION 5 — NET-SEC TRACKS
# ────────────────────────────────────────────────────────────────
add_page_break()
section_label("5 — Network Security AAS: Elective Track Structure (choose ONE, 12 credits)")

tracks = [
    {
        "icon": "🔌", "title": "Cisco-based Network Operations",
        "fill": SH_TEAL,
        "courses_2627": "CSCO 120  →  CSCO 121  →  CSCO 220 or CSCO 230",
        "courses_2728": "(same)",
        "notes": "9 credits (3 courses) + 1 open elective = 12 credits total.\nOld code CSCO 230B satisfies the CSCO 220/230 slot."
    },
    {
        "icon": "☁️", "title": "Cloud Security",
        "fill": SH_LIGHT,
        "courses_2627": "CIT 211  ·  CIT 213  ·  CIT 214  ·  CSEC 100",
        "courses_2728": "(same)",
        "notes": "4 courses = 12 credits. Fully prescribed."
    },
    {
        "icon": "🔍", "title": "Digital Forensics",
        "fill": SH_GREEN,
        "courses_2627": "CSEC 120  ·  CSEC 121  ·  CSEC 220  ·  CSEC 221",
        "courses_2728": "(same)",
        "notes": "4 courses × 3 credits = 12 credits."
    },
    {
        "icon": "📜", "title": "Compliance",
        "fill": SH_PURPLE,
        "courses_2627": "CSEC 101  ·  CSEC 191  ·  CSEC 225  ·  CSEC 226B",
        "courses_2728": "(same)",
        "notes": "4 courses = 12 credits. Fully prescribed."
    },
    {
        "icon": "⚔️", "title": "Offensive Security",
        "fill": SH_RED,
        "courses_2627": "CIT 274  ·  CSEC 181  ·  CSEC 281  +  1 open elective",
        "courses_2728": "CIT 274  ·  CSEC 181  ·  CSEC 281  ·  CSEC 282",
        "notes": "⚠️ CSEC 281 requires CSEC 181 completed with C or higher.\n"
                 "Old code CSEC 281B = CIT 274 (NOT the same as CSEC 281).\n"
                 "2027–28: CSEC 282 added, open elective removed."
    },
    {
        "icon": "🤖", "title": "Artificial Intelligence",
        "fill": SH_YELLOW,
        "courses_2627": "CIT 164  ·  CIT 223  ·  CIT 224  +  1 open elective",
        "courses_2728": "CIT 164  ·  CIT 223  ·  CIT 224  ·  CSEC 104  ·  CSEC 282",
        "notes": "2027–28: 5-course track (15 credits). Old CSEC 104B = CSEC 104."
    },
]

t = doc.add_table(rows=len(tracks)+1, cols=4)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.LEFT

# Header
for ci, lbl in enumerate(["Track", "2026–27 Courses", "2027–28 Courses", "Notes"]):
    cell = t.cell(0, ci)
    set_cell_bg(cell, "1e3a5f")
    run(cell.paragraphs[0], lbl, bold=True, color=RGBColor(0xe2,0xe8,0xf0), size=8.5)
    cell_borders(cell, "FFFFFF", 6)

for ri, track in enumerate(tracks):
    row = t.rows[ri+1]
    vals = [
        f"{track['icon']} {track['title']}",
        track["courses_2627"],
        track["courses_2728"],
        track["notes"],
    ]
    fills_tr = [track["fill"], "FFFFFF", "FFFFFF", "FAFAFA"]
    for ci, val in enumerate(vals):
        cell = row.cells[ci]
        set_cell_bg(cell, fills_tr[ci])
        p = cell.paragraphs[0]
        bold = ci == 0
        color = C_SUBHEAD if ci == 0 else (C_BLUE if ci in (1,2) else C_MUTED)
        run(p, val, bold=bold, color=color, size=8.5)
        cell_borders(cell, "DDDDDD", 4)

trk_widths = [Inches(1.6), Inches(2.2), Inches(2.2), Inches(2.5)]
for row in t.rows:
    for i, cell in enumerate(row.cells):
        cell.width = trk_widths[i]

doc.add_paragraph()
note_box(
    "Track selection happens AFTER all 8 core courses are completed. "
    "CSEC 286 (Applied Network Defense) is always the final core course.",
    fill=SH_LIGHT, icon="ℹ️", size=8.5
)


# ────────────────────────────────────────────────────────────────
#  SECTION 6 — PENDING ITEMS
# ────────────────────────────────────────────────────────────────
hline()
section_label("6 — Pending / Still To Implement")

pending = [
    ("1️⃣", "CSEC 290B Separate Handling (app.js)",
     "Currently grouped with the 112B/116B/212B deactivated courses. "
     "Needs its own urgent-enrollment message: 'Last offered Fall 2026 — register immediately.' "
     "Different from the CHFI substitution path.",
     SH_ORANGE),
    ("2️⃣", "CIT 263 'Does Not Transfer' Note (data.js)",
     "The AAS transfersToNew map is missing a CIT 263 entry. The CA table already shows ❌. "
     "Both should make it explicit that CIT 263 does not carry to either new map.",
     SH_LIGHT),
    ("3️⃣", "2027–28 Compare Button (app.js)",
     "The compare button currently only offers the 2026–27 map as a comparison target. "
     "DF AAS and Net-Sec old-map students should also be able to compare against 2027–28.",
     SH_PURPLE),
    ("4️⃣", "Compliance AAS Course Data (data.js)",
     "All Compliance AAS data is placeholder. No requirements have been entered yet. "
     "Blocked — awaiting instructor input.",
     SH_RED),
]

for icon_p, title_p, body_p, fill_p in pending:
    p = doc.add_paragraph()
    run(p, f"{icon_p}  {title_p}", bold=True, color=C_SUBHEAD, size=10.5)
    pPr = p._p.get_or_add_pPr()
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), "160")
    pSp.set(qn("w:after"),  "30")
    pPr.append(pSp)
    note_box(body_p, fill=fill_p, size=9)


# ────────────────────────────────────────────────────────────────
#  SAVE
# ────────────────────────────────────────────────────────────────
out = "/Users/lilmnstr/Documents/GitHub/degree-advisor/CIT_CyberSecurity_CourseAdvisor_Mindmap.docx"
doc.save(out)
print(f"Saved → {out}")
